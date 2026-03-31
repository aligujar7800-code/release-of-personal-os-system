"""
Text extractors for various file formats plus OCR.
"""

import os
import re
import logging

logger = logging.getLogger('MemoryOS-AI.extractors')


# ============================================================
# ✅ TEXT CLEANING HELPER
# ============================================================
def clean_text(text):
    """Extra spaces aur blank lines remove karo."""
    if not text:
        return ''
    # Multiple spaces ko single space karo
    text = re.sub(r' +', ' ', text)
    # 2 se zyada newlines ko 2 tak limit karo
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ============================================================
# ✅ PDF EXTRACTION (Same + cleaning added)
# ============================================================
def extract_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return clean_text('\n'.join(text_parts))
    except Exception as e:
        logger.error(f'PDF extraction error for {file_path}: {e}')
        return ''


# ============================================================
# ✅ DOCX EXTRACTION (Tables bhi add kiye)
# ============================================================
def extract_docx(file_path):
    """Extract text from a Word document including tables."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        # Normal paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # ✅ NEW - Tables ka text bhi nikalo
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return clean_text('\n'.join(text_parts))
    except Exception as e:
        logger.error(f'DOCX extraction error for {file_path}: {e}')
        return ''


# ============================================================
# ✅ TEXT FILE (Same + cleaning)
# ============================================================
def extract_text_file(file_path):
    """Extract text from plain text, markdown, or code files."""
    try:
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return clean_text(f.read())
            except (UnicodeDecodeError, UnicodeError):
                continue
        return ''
    except Exception as e:
        logger.error(f'Text extraction error for {file_path}: {e}')
        return ''


# ============================================================
# ✅ OCR (Same + image enhancement added)
# ============================================================
def ocr_image(image_path):
    """Extract text from an image using Tesseract OCR."""
    try:
        import pytesseract
        from PIL import Image, ImageFilter, ImageEnhance

        # Windows Tesseract path auto-detect
        if os.name == 'nt':
            common_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                os.path.join(
                    os.environ.get('LOCALAPPDATA', ''),
                    'Programs', 'Tesseract-OCR', 'tesseract.exe'
                ),
            ]
            for tess_path in common_paths:
                if os.path.exists(tess_path):
                    pytesseract.pytesseract.tesseract_cmd = tess_path
                    break

        img = Image.open(image_path)

        # ✅ NEW - Grayscale karo (OCR better hota hai)
        img = img.convert('L')

        # ✅ NEW - Contrast improve karo
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)

        # Resize if too large
        max_dim = 2000
        if max(img.size) > max_dim:
            ratio = max_dim / max(img.size)
            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
            img = img.resize(new_size, Image.LANCZOS)

        text = pytesseract.image_to_string(img)
        return clean_text(text)
    except Exception as e:
        logger.error(f'OCR error for {image_path}: {e}')
        return ''


# ============================================================
# ✅ NEW - EXCEL/CSV SUPPORT
# ============================================================
def extract_excel(file_path):
    """Extract text from Excel files."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text_parts.append(f"Sheet: {sheet}")
            for row in ws.iter_rows(values_only=True):
                row_text = ' | '.join(str(c) for c in row if c is not None)
                if row_text.strip():
                    text_parts.append(row_text)
        return clean_text('\n'.join(text_parts))
    except Exception as e:
        logger.error(f'Excel extraction error for {file_path}: {e}')
        return ''


def extract_csv(file_path):
    """Extract text from CSV files."""
    try:
        import csv
        text_parts = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                row_text = ' | '.join(cell for cell in row if cell.strip())
                if row_text:
                    text_parts.append(row_text)
        return clean_text('\n'.join(text_parts))
    except Exception as e:
        logger.error(f'CSV extraction error for {file_path}: {e}')
        return ''


# ============================================================
# ✅ NEW - AUTO DETECT - File type dekh ke sahi function call
# ============================================================
def extract_file(file_path):
    """
    File path do - automatically sahi extractor use karega.
    Tumhare project mein yeh main function use karo.
    """
    if not os.path.exists(file_path):
        logger.warning(f'File nahi mili: {file_path}')
        return ''

    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
        '.pdf':  extract_pdf,
        '.docx': extract_docx,
        '.doc':  extract_docx,
        '.txt':  extract_text_file,
        '.md':   extract_text_file,
        '.py':   extract_text_file,
        '.js':   extract_text_file,
        '.ts':   extract_text_file,
        '.html': extract_text_file,
        '.csv':  extract_csv,
        '.xlsx': extract_excel,
        '.xls':  extract_excel,
        '.png':  ocr_image,
        '.jpg':  ocr_image,
        '.jpeg': ocr_image,
        '.bmp':  ocr_image,
    }

    extractor = extractors.get(ext)
    if extractor:
        logger.info(f'Extracting: {file_path} ({ext})')
        return extractor(file_path)
    else:
        logger.warning(f'Unsupported file type: {ext}')
        return ''


# ============================================================
# ✅ TEST
# ============================================================
if __name__ == "__main__":
    # Test karo - apni koi file ka path do
    test_files = [
        "test.txt",
        "test.pdf",
        "test.docx",
    ]
    for f in test_files:
        if os.path.exists(f):
            result = extract_file(f)
            print(f"\n📄 {f}:")
            print(result[:200])  # Pehle 200 characters
        else:
            print(f"⚠️ File nahi mili: {f}")